from docx import Document
from sys import stdin
import pymorphy2


morph = pymorphy2.MorphAnalyzer()
data = list(map(lambda x: x.strip(), stdin.readlines()))

place = data.pop(0)
time = data.pop(0)


while data:
    name = data.pop(0)
    document = Document()
    document.add_heading('\tПриглашение на праздник ❤', 0).italic = True

    doc = document.add_paragraph("Сегодня – самый прекрасный и нежный, самый приятный и чудесный")
    doc.add_run(", самый красивый и загадочный праздник — Международный женский день, 8 Марта!")

    doc = document.add_paragraph(f"Поздравляем Вас, {name}, с этим весенним, неповторимым и ")
    doc.add_run("сказочным праздником! И приглашаем на мероприятие!")

    doc = document.add_heading(f"\tПразднование состоится {place.lower()}")
    doc.add_run(f" {time.lower()}").bold

    name = ' '.join(list(map(lambda x: morph.parse(x)[0].inflect({'gent'}).word.capitalize(),
                             name.split())))
    document.save(f"Приглашение для {name}.docx")


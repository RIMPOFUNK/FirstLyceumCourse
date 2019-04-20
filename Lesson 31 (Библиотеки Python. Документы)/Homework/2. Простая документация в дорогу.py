from docx import Document


def markdown_to_docx(text):
    docx = Document()

    docx.add_heading(text.split('\n')[0], 0)

    for i in text.split('\n'):
        if i == '\n':
            docx.add_paragraph()
            continue
        if i.startswith('#'):
            level = len(i.split()[0])
            docx.add_heading(f"{' '.join(i.split()[1:])}", level=level)

        if i.startswith('*') or i.startswith('-') or i.startswith('+'):
            docx.add_paragraph(f"{' '.join(i.split()[1:])}",
                       style='List Bullet')

        if any(i.startswith(j + '.') for j in "1234567890"):
            docx.add_paragraph(f"{' '.join(i.split()[1:])}",
                       style='List Number')
        docx.add_paragraph('', style="Intense Quote")




    docx.save('res.docx')

markdown_to_docx("""test01
Абзацы создаются при помощи пустой строки. Если вокруг текста сверху и снизу есть пустые строки, то текст превращается в абзац.

Чтобы сделать перенос строки вместо абзаца,
нужно поставить два пробела в конце предыдущей строки.

Заголовки отмечаются диезом `#` в начале строки, от одного до шести. Например:

# Заголовок первого уровня #
## Заголовок h2
### Заголовок h3
#### Заголовок h4
##### Заголовок h5
###### Заголовок h6

В декоративных целях заголовки можно «закрывать» с обратной стороны.

### Списки

Для разметки неупорядоченных списков можно использовать или `*`, или `-`, или `+`:

- элемент 1
- элемент 2
- элемент ...

Упорядоченный список:

1. элемент 1
2. элемент 2
3. элемент 3
4. Donec sit amet nisl. Aliquam semper ipsum sit amet velit. Suspendisse id sem consectetuer libero luctus adipiscing.

На самом деле не важно как в коде пронумерованы пункты, главное, чтобы перед элементом списка стояла цифра (любая) с точкой. Можно сделать и так:

0. элемент 1
0. элемент 2
0. элемент 3
0. элемент 4""")